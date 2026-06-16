from __future__ import annotations

from pathlib import Path
import subprocess
import sys

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SRC = ROOT / "final_materials"
OUT = ROOT / "out"


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for style_name, size, color in [
        ("Heading 1", 16, RGBColor(0x00, 0x00, 0x00)),
        ("Heading 2", 13, RGBColor(0x00, 0x00, 0x00)),
    ]:
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(12)
        style.paragraph_format.space_after = Pt(6)


def add_markdown(doc: Document, text: str) -> None:
    paragraph_lines: list[str] = []

    def flush_paragraph() -> None:
        if not paragraph_lines:
            return
        paragraph = doc.add_paragraph(" ".join(paragraph_lines))
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        paragraph_lines.clear()

    for raw_line in text.splitlines():
        line = raw_line.strip()
        if not line:
            flush_paragraph()
            continue
        if line.startswith("# "):
            flush_paragraph()
            doc.add_heading(line[2:], level=1)
            continue
        if line.startswith("## "):
            flush_paragraph()
            doc.add_heading(line[3:], level=2)
            continue
        paragraph_lines.append(line)
    flush_paragraph()


def build_docx(markdown_path: Path, docx_path: Path) -> None:
    doc = Document()
    configure_document(doc)
    add_markdown(doc, markdown_path.read_text(encoding="utf-8"))
    doc.save(docx_path)


def convert_docx_to_pdf(docx_path: Path, pdf_path: Path) -> None:
    OUT.mkdir(parents=True, exist_ok=True)
    subprocess.run(
        [
            "soffice",
            "--headless",
            "--convert-to",
            "pdf",
            "--outdir",
            str(OUT),
            str(docx_path),
        ],
        check=True,
    )
    generated = OUT / f"{docx_path.stem}.pdf"
    if generated != pdf_path:
        generated.replace(pdf_path)


def main() -> int:
    OUT.mkdir(parents=True, exist_ok=True)
    for stem in [
        "cover_letter_final_materials",
        "responses_to_reviewers_final_materials",
    ]:
        markdown_path = SRC / f"{stem}.md"
        docx_path = OUT / f"{stem}.docx"
        pdf_path = OUT / f"{stem}.pdf"
        build_docx(markdown_path, docx_path)
        convert_docx_to_pdf(docx_path, pdf_path)
        print(f"wrote {docx_path.relative_to(ROOT)}")
        print(f"wrote {pdf_path.relative_to(ROOT)}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
